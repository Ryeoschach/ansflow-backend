import os
import json
import logging
import requests
from django.conf import settings
from typing import List, Optional
from pydantic import ConfigDict
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_community.retrievers import BM25Retriever

# 核心导入修复：恢复环境兼容性导入
try:
    from langchain.retrievers import EnsembleRetriever
    from langchain.retrievers.contextual_compression import ContextualCompressionRetriever
except ImportError:
    try:
        from langchain_classic.retrievers import EnsembleRetriever
        from langchain_classic.retrievers.contextual_compression import ContextualCompressionRetriever
    except ImportError:
        # 最后的兜底
        EnsembleRetriever = None
        ContextualCompressionRetriever = None

from langchain_core.documents import Document
from langchain_core.documents.compressor import BaseDocumentCompressor
import jieba

# 必须在导入任何 AI 库之前设置环境变量，确保缓存路径被全局识别
CACHE_DIR = os.path.join(settings.BASE_DIR, ".model_cache")
if not os.path.exists(CACHE_DIR):
    os.makedirs(CACHE_DIR, exist_ok=True)

os.environ["FLASHRANK_CACHE_DIR"] = CACHE_DIR
os.environ["FASTEMBED_CACHE_PATH"] = CACHE_DIR

logger = logging.getLogger(__name__)

def chinese_tokenizer(text: str):
    """使用 jieba 进行中文分词，提高 BM25 的匹配精度"""
    return jieba.lcut_for_search(text.lower())

class GenericRerank(BaseDocumentCompressor):
    """通用的远程重排序器，支持所有遵循 POST /rerank 规范的引擎 (如 Xinference)"""
    base_url: str
    model_name: str
    top_n: int = 5
    api_key: Optional[str] = None
    model_config = ConfigDict(arbitrary_types_allowed=True)

    def compress_documents(
        self,
        documents: List[Document],
        query: str,
        callbacks: Optional[any] = None,
    ) -> List[Document]:
        if not documents:
            return []
        try:
            base = self.base_url.rstrip('/')
            url = f"{base}/rerank" if "/v1" in base else f"{base}/v1/rerank"
            payload = {
                "model": self.model_name,
                "query": query,
                "documents": [doc.page_content for doc in documents],
                "top_n": self.top_n
            }
            headers = {"Authorization": f"Bearer {self.api_key}"} if self.api_key else {}
            response = requests.post(url, json=payload, headers=headers, timeout=(5, 25))
            response.raise_for_status()
            result = response.json()
            final_docs = []
            data = result.get("results") or result.get("data") or []
            for item in data[:self.top_n]:
                idx = item.get("index")
                score = item.get("relevance_score") or item.get("score")
                if idx is not None and idx < len(documents):
                    doc = documents[idx]
                    doc.metadata["rerank_score"] = score
                    final_docs.append(doc)
            return final_docs if final_docs else documents[:self.top_n]
        except Exception as e:
            logger.error(f"Rerank failed: {e}")
            return documents[:self.top_n]

from .models import AIConfig, AIModel, AIProvider, KnowledgeChunk, KnowledgeDocument, KnowledgeBase
from .vision_parser import VisionParser

class RAGService:
    PERSONALITIES = {
        'professional': {
            'prefix': '你是一个资深的 AnsFlow SRE 专家。你的回答应该专业、客观，包含必要的技术细节和代码示例。'
        },
        'concise': {
            'prefix': '你是一个高效的运维助手。请用最简短的语言回答问题，直接给出结论和命令，避免废话。'
        },
        'humorous': {
            'prefix': '你是一个热爱生活的运维老工。虽然运维很苦，但你的回答总能带着一点点幽默感，偶尔调侃一下 Bug。'
        }
    }

    _vectorstore_cache = {}
    _embeddings_cache = {}
    _reranker_cache = {}

    def __init__(self, collection_name: str = "ansflow_docs", personality: str = 'professional',
                 llm_id: int = None, embedding_id: int = None):
        self.personality = self.PERSONALITIES.get(personality, self.PERSONALITIES['professional'])
        self.persist_directory = os.path.join(settings.BASE_DIR, "chroma_db")
        self.cache_directory = CACHE_DIR
        
        self.config = AIConfig.objects.filter(name="default").first()
        self.llm_config = self._get_model_config(llm_id, "llm")
        self.emb_config = self._get_model_config(embedding_id, "embedding")
        self.rerank_config = self._get_model_config(None, "rerank")

        emb_key = f"{self.emb_config.get('name', 'default')}_{self.emb_config.get('provider_type', 'local')}"
        if emb_key not in self._embeddings_cache:
            self._embeddings_cache[emb_key] = self._init_embeddings()
        self.embeddings = self._embeddings_cache[emb_key]

        rerank_key = f"{self.rerank_config.get('name', 'none')}_{self.rerank_config.get('provider_type', 'none')}"
        if rerank_key not in self._reranker_cache:
            rerank_ptype = self.rerank_config.get('provider_type')
            reranker_base = self.rerank_config.get('base_url')
            reranker_model = self.rerank_config.get('name')
            is_remote_configured = rerank_ptype != "local" and reranker_base
            new_reranker = None
            if is_remote_configured:
                try:
                    new_reranker = GenericRerank(base_url=reranker_base, model_name=reranker_model, api_key=self.rerank_config.get('api_key'))
                except: pass
            if new_reranker is None and not is_remote_configured:
                try:
                    from langchain_community.document_compressors.flashrank_rerank import FlashrankRerank
                    from flashrank import Ranker
                    target_model = reranker_model if reranker_model else "ms-marco-MiniLM-L-12-v2"
                    flashrank_client = Ranker(model_name=target_model, cache_dir=CACHE_DIR)
                    new_reranker = FlashrankRerank(client=flashrank_client, model=target_model)
                except: pass
            self._reranker_cache[rerank_key] = new_reranker
        self.reranker = self._reranker_cache[rerank_key]
        
        safe_model_name = self.emb_config.get('name', 'default').replace("/", "_").replace("-", "_")
        full_collection_name = f"{collection_name}_{safe_model_name}"
        if full_collection_name not in self._vectorstore_cache:
            self._vectorstore_cache[full_collection_name] = Chroma(collection_name=full_collection_name, embedding_function=self.embeddings, persist_directory=self.persist_directory)
        self.vectorstore = self._vectorstore_cache[full_collection_name]
        self.llm = self._init_llm()

    def _get_model_config(self, model_id: int, model_type: str):
        model = None
        if model_id:
            model = AIModel.objects.filter(id=model_id, model_type=model_type).first()
        if not model:
            global_config = AIConfig.objects.filter(name="default").first()
            if global_config:
                if model_type == "llm": model = global_config.default_llm
                elif model_type == "embedding": model = global_config.default_embedding
                elif model_type == "rerank": model = global_config.default_rerank
        if model:
            return {"name": model.name, "provider_type": model.provider.provider_type, "base_url": model.provider.base_url, "api_key": model.provider.get_decrypted_key()}
        return {}

    def _init_embeddings(self):
        ptype = self.emb_config.get('provider_type', 'local')
        name = self.emb_config.get('name', 'default')
        if ptype == "local" or "BAAI" in name:
            from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
            return FastEmbedEmbeddings(model_name=name, cache_dir=self.cache_directory)
        
        from langchain_openai import OpenAIEmbeddings
        base_url = self.emb_config.get('base_url', '').rstrip('/')
        # 自动补全 /v1，适配本地常用服务
        if base_url and not base_url.endswith('/v1'):
            base_url = f"{base_url}/v1"
            
        return OpenAIEmbeddings(
            model=name, 
            api_key=self.emb_config.get('api_key') or "not-needed", 
            base_url=base_url,
            chunk_size=16 # 降低批处理大小，适配本地低配环境
        )

    def _init_llm(self):
        return ChatOpenAI(model=self.llm_config['name'], api_key=self.llm_config['api_key'], base_url=self.llm_config['base_url'], streaming=True)

    def get_retriever(self, kb_id: int = None):
        from .models import KnowledgeChunk
        top_k = self.config.rag_top_k if self.config else 3
        vector_retriever = self.vectorstore.as_retriever(search_kwargs={"k": top_k * 4})
        
        chunks_qs = KnowledgeChunk.objects.filter(is_active=True)
        if kb_id: chunks_qs = chunks_qs.filter(document__kb_id=kb_id)
        
        chunks = list(chunks_qs.values('content', 'metadata'))
        if not chunks:
            if self.reranker and ContextualCompressionRetriever:
                self.reranker.top_n = top_k
                return ContextualCompressionRetriever(base_compressor=self.reranker, base_retriever=vector_retriever)
            vector_retriever.search_kwargs["k"] = top_k
            return vector_retriever
        bm25_retriever = BM25Retriever.from_texts(texts=[c['content'] for c in chunks], metadatas=[c['metadata'] for c in chunks], preprocess_func=chinese_tokenizer)
        bm25_retriever.k = top_k * 4
        
        if EnsembleRetriever:
            ensemble_retriever = EnsembleRetriever(retrievers=[vector_retriever, bm25_retriever], weights=[self.config.rag_vector_weight if self.config else 0.7, self.config.rag_bm25_weight if self.config else 0.3])
            if self.reranker and ContextualCompressionRetriever:
                self.reranker.top_n = top_k
                return ContextualCompressionRetriever(base_compressor=self.reranker, base_retriever=ensemble_retriever)
            return ensemble_retriever
        return vector_retriever

    def get_chat_chain(self, history_id: int = None, auth_context: dict = None):
        kbs = KnowledgeBase.objects.all()
        kb_catalog = "\n".join([f"- {kb.name}: {kb.description or '无描述'} (文档数: {kb.documents.count()})" for kb in kbs])
        chat_memory_str = ""
        if history_id:
            from .models import AIChatMessage
            prev_messages = AIChatMessage.objects.filter(history_id=history_id).order_by('create_time')[:10]
            for m in prev_messages:
                chat_memory_str += f"{'用户' if m.role == 'user' else '助手'}: {m.content}\n"
        template = """{prefix}
【系统知识目录】
你连接了 AnsFlow 的全量知识库系统，当前包含以下库：
{kb_catalog}

【特殊指令】
1. 资产编排：写剧本输出 `__ANSIBLE_DRAFT__: {{"name": "...", "content": "..."}}`。
2. 流水线编排：输出 `__PIPELINE_DRAFT__: {{"nodes": [...], "edges": [...]}}`。

参考内容：
{context}

对话历史：
{chat_history}

用户问题：{question}
你的回答："""
        prompt = ChatPromptTemplate.from_template(template)
        def context_retriever(query):
            docs = self.retrieve_with_threshold(query, kb_id=None)
            return "\n\n".join(doc.page_content for doc in docs)
        chain = ({"context": context_retriever, "question": RunnablePassthrough(), "prefix": lambda x: self.personality['prefix'], "kb_catalog": lambda x: kb_catalog, "chat_history": lambda x: chat_memory_str} | prompt | self.llm | StrOutputParser())
        return chain

    def chat_stream(self, question: str, history_id: int = None, auth_context: dict = None):
        chain = self.get_chat_chain(history_id=history_id, auth_context=auth_context)
        yield from chain.stream(question)

    def diagnose_log(self, log_content: str, context_info: dict, auth_context: dict = None):
        kbs = KnowledgeBase.objects.all()
        kb_catalog = "\n".join([f"- {kb.name}: {kb.description or '无描述'}" for kb in kbs])
        template = """{prefix}
作为专业 SRE，分析以下日志并给出诊断结论。
【系统知识目录】
{kb_catalog}
【执行上下文】
- 类型: {target_type}, 名称: {target_name}, 摘要: {error_summary}
【错误日志】
{log_content}
【参考知识库】
{context}
请给出：故障根因、修复建议（包含编排指令标记）、预防措施。
"""
        prompt = ChatPromptTemplate.from_template(template)
        search_text = f"诊断 {context_info.get('name')} 错误: {context_info.get('summary')}"
        referenced_docs = self.retrieve_with_threshold(search_text, kb_id=None)
        refs = [{"id": d.metadata.get('document_id'), "title": d.metadata.get('title') or "参考文档"} for d in referenced_docs if d.metadata.get('document_id')]
        if refs: yield f"__REFERENCES__:{json.dumps(refs)}\n"
        chain = ({"context": lambda x: self.format_docs(referenced_docs), "target_type": lambda x: x["target_type"], "target_name": lambda x: x["target_name"], "error_summary": lambda x: x["error_summary"], "log_content": lambda x: x["log_content"], "prefix": lambda x: self.personality['prefix'], "kb_catalog": lambda x: kb_catalog} | prompt | self.llm | StrOutputParser())
        yield from chain.stream({"log_content": log_content, "target_type": context_info.get("type", "Unknown"), "target_name": context_info.get("name", "Unknown"), "error_summary": context_info.get("summary", "failed")})

    def diagnose_alert(self, query: str):
        """
        同步调用 AI 诊断告警
        """
        referenced_docs = self.retrieve_with_threshold(query, kb_id=None)
        template = """{prefix}
你是一个资深的 SRE 专家。请针对以下告警信息进行深度诊断。
【参考知识库】
{context}
【告警详情】
{query}
请给出：
1. 故障根因分析
2. 修复建议（包括具体的命令或操作步骤）
3. 预防措施
"""
        prompt = ChatPromptTemplate.from_template(template)
        # 封装同步执行链
        chain = (
            {"context": lambda x: self.format_docs(referenced_docs), "query": RunnablePassthrough(), "prefix": lambda x: self.personality['prefix']}
            | prompt
            | self.llm
            | StrOutputParser()
        )
        return chain.invoke(query)

    def retrieve_with_threshold(self, query: str, kb_id: int = None):
        retriever = self.get_retriever(kb_id=kb_id)
        docs = retriever.invoke(query)
        threshold = self.config.rag_score_threshold if self.config else 0.0
        if threshold <= 0: return docs
        return [d for d in docs if d.metadata.get('rerank_score', 1.0) >= threshold]

    def format_docs(self, docs):
        return "\n\n".join(doc.page_content for doc in docs)

    def add_knowledge(self, content: str, metadata: dict = None, kb_id: int = None, title: str = None):
        kb = KnowledgeBase.objects.filter(id=kb_id).first() if kb_id else KnowledgeBase.objects.get_or_create(name="默认知识库", defaults={"collection_name": "ansflow_docs"})[0]
        kd = KnowledgeDocument.objects.create(kb=kb, title=title or f"AI Export {os.urandom(2).hex()}", content=content, source_type="ai_export", metadata=metadata or {})
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=120)
        splits = text_splitter.split_text(content)
        ids = [f"{kd.id}_{i}" for i in range(len(splits))]
        self.vectorstore.add_texts(texts=splits, ids=ids, metadatas=[{"document_id": kd.id, "kb_id": kb.id} for _ in splits])
        return True

    def ingest_document(self, file_path: str, kb_id: int = None, document_id: int = None):
        """
        升级版摄取流程：
        1. 路由分类 (根据用户选择的模式)
        2. 专有解析提取 (注入 Prompt)
        3. 文本清洗与标准化
        4. 智能切片 (num_ctx 感知)
        5. 向量化入库
        """
        from .models import KnowledgeDocument, KnowledgeBase, KnowledgeChunk
        kb = KnowledgeBase.objects.filter(id=kb_id).first() if kb_id else KnowledgeBase.objects.get_or_create(name="默认知识库")[0]
        
        kd = None
        if document_id:
            kd = KnowledgeDocument.objects.get(id=document_id)
        
        parser_type = kd.parser_type if kd else "auto"
        parsing_prompt = kd.parsing_prompt if kd else None
        
        logger.info(f"[Ingest] Starting ingestion for {file_path}. Mode: {parser_type}, Prompt: {parsing_prompt[:50] if parsing_prompt else 'None'}")
        
        # 1. & 2. 解析提取
        if kd:
            kd.status = 'parsing'
            kd.save(update_fields=['status'])
            
        full_content = self._extract_content(file_path, parser_type, parsing_prompt)
        
        if not full_content:
            logger.warning(f"[Ingest] No content extracted from {file_path}")
            if kd:
                kd.status = 'error'
                kd.save(update_fields=['status'])
            return 0

        # 3. 文本清洗
        if kd:
            kd.status = 'cleaning'
            kd.save(update_fields=['status'])
        full_content = self._clean_text(full_content)

        # 获取或更新文档记录
        if not kd:
            kd, _ = KnowledgeDocument.objects.get_or_create(
                kb=kb, 
                title=os.path.basename(file_path), 
                defaults={"content": full_content, "file_path": file_path, "source_type": "file", "status": "chunking"}
            )
        else:
            kd.content = full_content
            kd.status = 'chunking'
            kd.save()
        
        # 4. 智能切片 (计算 num_ctx)
        chunk_size, chunk_overlap = self._get_optimal_chunk_params()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        splits = text_splitter.split_text(full_content)
        
        # 5. 向量化入库
        kd.status = 'indexing'
        kd.save(update_fields=['status'])
        
        ids = [f"{kd.id}_{i}" for i in range(len(splits))]
        self.vectorstore.add_texts(
            texts=splits, 
            ids=ids, 
            metadatas=[{
                'document_id': kd.id, 
                'kb_id': kb.id, 
                'chunk_index': i,
                'source': kd.title,
                'ingest_mode': parser_type,
                'type': 'human_verified_knowledge' if parser_type != 'auto' else 'raw_knowledge'
            } for i in range(len(splits))]
        )
        
        chunk_objs = [KnowledgeChunk(
            document=kd, 
            content=s, 
            vector_id=ids[i], 
            index=i, 
            metadata={
                'document_id': kd.id, 
                'kb_id': kb.id,
                'source': kd.title
            }
        ) for i, s in enumerate(splits)]
        
        KnowledgeChunk.objects.filter(document=kd).delete()
        KnowledgeChunk.objects.bulk_create(chunk_objs)
        
        kd.status, kd.chunk_count = "ready", len(splits)
        kd.save()
        return len(splits)

    def _extract_content(self, file_path: str, mode: str, prompt: Optional[str] = None) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        
        # 如果是自动模式，根据后缀路由
        if mode == "auto":
            if ext in ['.txt', '.md', '.json', '.yaml', '.yml']:
                mode = "native"
            elif ext in ['.png', '.jpg', '.jpeg']:
                mode = "ocr"
            elif ext in ['.pdf', '.docx']:
                mode = "hybrid"
            else:
                mode = "native"

        logger.info(f"[Extract] Final mode selected: {mode} for extension {ext}")

        if mode == "native":
            try:
                if ext in ['.pdf']:
                    import fitz
                    doc = fitz.open(file_path)
                    return "\n".join([page.get_text() for page in doc])
                from langchain_community.document_loaders import TextLoader
                loader = TextLoader(file_path, encoding='utf-8')
                return "\n".join([d.page_content for d in loader.load()])
            except Exception as e:
                logger.error(f"[Ingest] Native extraction failed: {e}")
                return ""

        elif mode == "ocr":
            # 强制走 OCR 解析
            return self._ocr_extraction(file_path, ext, prompt)

        elif mode == "hybrid":
            # 混合模式：优先尝试复杂解析
            if ext == '.docx':
                return self._ocr_extraction(file_path, ext, prompt)
            elif ext == '.pdf':
                import fitz
                doc = fitz.open(file_path)
                text_parts = [page.get_text() for page in doc]
                full_text = "\n".join(text_parts)
                # 判据优化：如果文本极少或用户提供了提示词，则尝试 OCR 增强
                if len(full_text.strip()) < 150 or (prompt and len(prompt.strip()) > 0):
                    logger.info(f"[Extract] Hybrid PDF triggering OCR (text_len: {len(full_text.strip())}, has_prompt: True)")
                    return self._ocr_extraction(file_path, ext, prompt)
                return full_text
            
        return ""

    def _clean_text(self, text: str) -> str:
        """文本清洗与标准化"""
        import re
        # 1. 统一换行符
        text = text.replace('\r\n', '\n')
        # 2. 合并 OCR 可能产生的断行 (非段落断行)
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
        # 3. 去除重复的空白字符
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 4. 去除常见的 OCR 杂质 (如只有 1-2 个字符的行)
        lines = [line.strip() for line in text.split('\n') if len(line.strip()) > 1]
        return "\n".join(lines)

    def _get_optimal_chunk_params(self) -> tuple:
        """根据模型 num_ctx 计算最优切片"""
        num_ctx = 4096
        if self.llm_config and 'name' in self.llm_config:
            # 尝试从模型记录中获取
            model = AIModel.objects.filter(name=self.llm_config['name']).first()
            if model:
                num_ctx = model.num_ctx
        
        # 策略：分块大小建议为窗口的 15-20%
        # 向上取整到 100 的倍数，范围在 400-1200 之间
        ideal_size = min(max(int(num_ctx * 0.2), 400), 1200)
        overlap = int(ideal_size * 0.15)
        return ideal_size, overlap

    def _ocr_extraction(self, file_path: str, ext: str, prompt: Optional[str] = None) -> str:
        """内部视觉解析流程"""
        # 优先使用全局配置中指定的视觉模型
        v_model_id = self.config.default_vision.id if self.config and self.config.default_vision else None
        v_parser = VisionParser(model_id=v_model_id)
        all_text = []
        try:
            if ext == '.pdf':
                import fitz
                doc = fitz.open(file_path)
                for page in doc:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")
                    page_text = v_parser.parse_image(img_bytes, custom_prompt=prompt)
                    all_text.append(f"--- Page {page.number + 1} ---\n{page_text}")
                doc.close()
            elif ext in ['.png', '.jpg', '.jpeg']:
                with open(file_path, 'rb') as f:
                    all_text.append(v_parser.parse_image(f.read(), custom_prompt=prompt))
            elif ext == '.docx':
                # Word 目前仍通过 python-docx 提取文本，后续可扩展视觉解析
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                for para in doc.paragraphs:
                    if para.text.strip(): all_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        all_text.append(" | ".join([cell.text.strip() for cell in row.cells]))
        except Exception as e:
            logger.error(f"[Vision] OCR failed for {file_path}: {e}")
        return "\n\n".join(all_text)

    def generate_dag(self, prompt_text: str, context_data: dict = None):
        template = "你是一个专业的流水线专家。生成 JSON 格式组织好的 DAG：{prompt_text}"
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | self.llm | StrOutputParser()
        return chain.invoke({"prompt_text": prompt_text})

    def explain_pipeline(self, nodes: list, edges: list):
        template = "解释以下流水线逻辑：{pipeline}"
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | self.llm | StrOutputParser()
        return chain.invoke({"pipeline": json.dumps({"nodes": nodes, "edges": edges})})

    def rewrite_query(self, query: str):
        return query

    def delete_document(self, document_id: int):
        from .models import KnowledgeChunk
        chunks = KnowledgeChunk.objects.filter(document_id=document_id)
        vids = list(chunks.values_list('vector_id', flat=True))
        if vids: self.vectorstore.delete(ids=vids)
        chunks.delete()
        return True

    def reindex_all(self, kb_id: int):
        """
        全量重建知识库索引：
        1. 清理向量库
        2. 重新处理所有文档
        """
        from .models import KnowledgeDocument, KnowledgeChunk
        docs = KnowledgeDocument.objects.filter(kb_id=kb_id)
        
        # 1. 批量获取并删除该 KB 的所有向量
        all_chunks = KnowledgeChunk.objects.filter(document__kb_id=kb_id)
        vids = list(all_chunks.values_list('vector_id', flat=True))
        if vids:
            try:
                self.vectorstore.delete(ids=vids)
            except Exception as e:
                logger.error(f"[Reindex] Failed to clear vectorstore: {e}")
        
        # 2. 逐个重新索引文档
        success_count = 0
        for doc in docs:
            try:
                if doc.source_type == 'file' and doc.file_path:
                    # 文件类：走完整摄取逻辑
                    self.ingest_document(doc.file_path, kb_id=kb_id, document_id=doc.id)
                else:
                    # 手动或 AI 导出类：直接使用 content 重新切片
                    doc.status = 'chunking'
                    doc.save(update_fields=['status'])
                    
                    # 使用与 ingest_document 相同的智能切片逻辑
                    chunk_size, chunk_overlap = self._get_optimal_chunk_params()
                    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
                    splits = text_splitter.split_text(doc.content)
                    
                    ids = [f"{doc.id}_{i}" for i in range(len(splits))]
                    self.vectorstore.add_texts(
                        texts=splits, 
                        ids=ids, 
                        metadatas=[{
                            'document_id': doc.id, 
                            'kb_id': kb_id, 
                            'chunk_index': i,
                            'source': doc.title
                        } for i in range(len(splits))]
                    )
                    
                    chunk_objs = [KnowledgeChunk(
                        document=doc, 
                        content=s, 
                        vector_id=ids[i], 
                        index=i, 
                        metadata={'document_id': doc.id, 'kb_id': kb_id}
                    ) for i, s in enumerate(splits)]
                    
                    KnowledgeChunk.objects.filter(document=doc).delete()
                    KnowledgeChunk.objects.bulk_create(chunk_objs)
                    
                    doc.status, doc.chunk_count = "ready", len(splits)
                    doc.save()
                    
                success_count += 1
            except Exception as e:
                logger.error(f"[Reindex] Failed to process doc {doc.id}: {e}")
                doc.status = 'error'
                doc.save(update_fields=['status'])
                
        return success_count

    def update_chunk(self, chunk_id: int, content: str):
        from .models import KnowledgeChunk
        chunk = KnowledgeChunk.objects.filter(id=chunk_id).first()
        if chunk:
            self.vectorstore.add_texts(texts=[content], ids=[chunk.vector_id], metadatas=[chunk.metadata])
            return True
        return False

    def delete_chunk(self, chunk_id: int):
        from .models import KnowledgeChunk
        chunk = KnowledgeChunk.objects.filter(id=chunk_id).first()
        if chunk and chunk.vector_id:
            self.vectorstore.delete(ids=[chunk.vector_id])
            chunk.delete()
            return True
        return False
